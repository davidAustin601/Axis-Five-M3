
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from openpyxl.styles import Alignment
import openpyxl
import docx
from docx.api import Document
import argparse
import datetime
import os
from shutil import copyfile
from datetime import datetime

from openpyxl.styles.numbers import NumberFormat

#PROGRAM ARGUMENTS [NAME, NUM PER MONTH, DATE, TIME]

#VARIABLES Defined
#...Program Arguments
argTaskType = " "
argExcel_Filename = " "
argWord_FileDIR = " "
argNumPerMonth = " "
argDateScreen = " "
argInitialContactTime = " "
argResponseTime = " "
argPerson = " "
argMonthYear = " "
argScreenshotDIR = " "
#...Excel Styles
globExcelGreen = "93D051"
globExcelBLUE = "00B0F0"
globExcelRed = "FE0000"
#...Excel Column Indices
intcolumn_NamePerson = 0
intcolumn_TestNum = 8
intcolumn_ScreenDate = 9
intcolumn_ContactDate = 5 
#...Misc Global Variables
globDropbox = '/Users/davidaustin/Dropbox/MARP/DRUG SCREENS/Logs/Monthly Reports/'


#Define Methods
#...Global Methods
def getContactLogPath():
    contactLogSavePath = argWord_FileDIR

    if argNumPerMonth == '1':
        contactLogSavePath += "MARP_Contact_Log.docx"
    elif argNumPerMonth == '2':
        contactLogSavePath += "MARP_Contact_Log 2.docx"
    else:
        contactLogSavePath = " "

    return contactLogSavePath
#...Time Maniuplation Methods
def time_converter(time):
    d = datetime.strptime(time, "%H:%M")
    return d.strftime("%I:%M %p")

#...Log a new contact date
def LogContactDate():

    try:

        #declare runLog as global
        tempLog = ""
        #variable for Seeing if a date has already been logged
        dateLogged = False

        #Create Workbook and Sheet
        workbook = load_workbook(filename=argExcel_Filename)
        sheet = workbook.active

        
        #******************LOOGGGING*************************
        #print("! (1/5) MARP Drugscreen Call List Loaded")
        tempLog += "! (1/5) MARP Drugscreen Call List Loaded" + "\n"

        #cycle through rows
        rowTicker = 0
        for row in sheet.iter_rows(values_only=True):

            #Skip First Row
            if rowTicker > 0:

                rowTicker += 1

                #current row Person
                curNamePerson = row[intcolumn_NamePerson]
                
                #compare current name to global name
                if argPerson.upper() == str(curNamePerson).upper():

                    #******************LOOGGGING*************************
                    #print("! (2/5) MARP Member Row Found")
                    tempLog += "! (2/5) MARP Member Row Found" + "\n"

                    #Check to see if cell is not populated
                    curCellValue = sheet.cell(row=rowTicker, column=intcolumn_ContactDate).value

                    if not curCellValue and not dateLogged:

                        #Concat Date and Time
                        curDateTime = str(argDateScreen) + " " + str(argResponseTime)
                        #Change Data and make sure justified right
                        sheet.cell(row=rowTicker, column=intcolumn_ContactDate).value = curDateTime
                        dateLogged = True
                        #******************LOOGGGING*************************
                        #print("! (3/5) Contact Date Logged")
                        tempLog += "! (3/5) Contact Date Logged" + "\n"

                        sheet.cell(row=rowTicker, column=intcolumn_ContactDate).alignment = Alignment(horizontal="right")

                        #Cycle through columns and change them to green
                        columnTicker = 0
                        for tempCell in sheet.iter_cols(values_only=True):
                            columnTicker += 1
                            sheet.cell(row=rowTicker, column=columnTicker).fill = PatternFill(fill_type='solid', start_color=globExcelGreen, end_color=globExcelGreen)

                        #******************LOOGGGING*************************
                        #print("! (4/5) MARP Member Row Formated to Green")
                        tempLog += "! (4/5) MARP Member Row Formated to Green" + "\n"

                        #save excel file
                        workbook.save(filename=argExcel_Filename)

                        #******************LOOGGGING*************************
                        #print("! (5/5) MARP Drugscreen Call List Saved")
                        tempLog += "! (5/5) MARP Drugscreen Call List Saved" + "\n"

                    else:
                        #******************LOOGGGING*************************
                        #print("* MARP Contact Date Already Populated")
                        tempLog += "* MARP Contact Date Already Populated" + "\n"
            else:
                rowTicker += 1
    except:
        #print("* Error Editing MARP Drugscreen Call List")
        tempLog += "* Error Editing MARP Drugscreen Call List" + "\n"

    return tempLog

#...Edit Contact Log Word Document
def docx_find_replace_text(doc, search_text, replace_text):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            # Replace strings and retain the same style.
            # The text to be replaced can be split over several runs so
            # search through, identify which runs need to have text replaced
            # then replace the text in those identified
            started = False
            search_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replace_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break

                if search_text[search_index] not in inline[i].text and not started:
                    # keep looking ...
                    continue

                # case 2: search for partial text, find first run
                if search_text[search_index] in inline[i].text and inline[i].text[-1] in search_text and not started:
                    # check sequence
                    start_index = inline[i].text.find(search_text[search_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != search_text[search_index]:
                            # no match so must be false positive
                            break
                    if search_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    search_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if search_index != len(search_text):
                        continue
                    else:
                        # found all chars in search_text
                        found_all = True
                        break

                # case 2: search for partial text, find subsequent run
                if search_text[search_index] in inline[i].text and started and not found_all:
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == search_text[search_index]:
                            search_index += 1
                            chars_found += 1
                        else:
                            break
                    # no match so must be end
                    found_runs.append((i, 0, chars_found))
                    if search_index == len(search_text):
                        found_all = True
                        break

            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], str(replace_text))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text
            # print(p.text)
def EditWordDocument():

    #create tempLog to return
    tempLog = ""

    try:

        #Split Name to get First Name
        splitName = argPerson.split()
        firstName = str(splitName[0])

        #Convert to 12 Hour
        convertedResponeTime = time_converter(argResponseTime)

        #define variables
        symbolName = "[PERSON]"
        symbolFirstName = "[FIRSTNAME]"
        symbolDate = "[DATE]"
        symbolInitialContactTime = "[INITIAL]"
        symbolResponseTime = "[RESPONSETIME]"
        
        #load word document=
        document = Document(globWordTemplate)

        

        #******************LOOGGGING*************************
        #print("! (1/3) MARP Contact Log Loaded")
        tempLog += "! (1/3) MARP Contact Log Loaded" + "\n"

        #Replace Name
        docx_find_replace_text(document, symbolName, argPerson)
        #Replace Fist Name
        docx_find_replace_text(document, symbolFirstName, firstName)
        #Replace Date
        docx_find_replace_text(document, symbolDate, argDateScreen)
        #Replace Initial Time
        docx_find_replace_text(document,symbolInitialContactTime, argInitialContactTime)
        #Replace Time
        docx_find_replace_text(document, symbolResponseTime, convertedResponeTime)
        
        #******************LOOGGGING*************************
        #print("! (2/3) MARP Contact Log Edited")
        tempLog += "! (2/3) MARP Contact Log Edited" + "\n"

        #Save Document
        contactLogSavePath = getContactLogPath()
        if not contactLogSavePath == " ":
            document.save(contactLogSavePath)

            #******************LOOGGGING*************************
            #print("! (3/3) MARP Contact Log Saved")
            tempLog += "! (3/3) MARP Contact Log Saved" + "\n"
        else:
            #print("* Error Editing MARP Contact Log")
            tempLog += "* Error Editing MARP Contact Log" + "\n"

    except:
        #print("* Error Editing MARP Contact Log")
        tempLog += "* Error Editing MARP Contact Log" + "\n"

    return tempLog

#...Update Dropbox Files
def CopyDocs():

    #create a tempLog to return
    tempLog = ""

    try:

        #Get Working Path
        dropboxPath = globDropbox + argMonthYear + "/" + argPerson + "/"

        #Get Contact Dropbox Destination Path
        dropboxWordPath = dropboxPath
    
        if argNumPerMonth == '1':
            dropboxWordPath += "MARP_Contact_Log.docx"
        elif argNumPerMonth == '2':
            dropboxWordPath += "MARP_Contact_Log 2.docx"
        else:
            dropboxWordPath = " "

        #Get Contact Log Source Path
        contactLogSourcePath = getContactLogPath()

       

        if not contactLogSourcePath == " " or not dropboxWordPath == " ":

            #Copy Word Doc to Dropbox
            copyfile(contactLogSourcePath, dropboxWordPath)
        
            #******************LOOGGGING*************************
            #print("! (1/3) MARP Contact Log Copied to Dropbox")
            tempLog += "! (1/3) MARP Contact Log Copied to Dropbox" + "\n"
        else:
            #******************LOOGGGING*************************
            #print("! * Error Copying MARP Contact Log to Dropbox")
            tempLog += "! * Error Copying MARP Contact Log to Dropbox"

        #Get Screenshot Filname and Path
        screenshotName = getScreenshotName()
        screenshotPath = argScreenshotDIR + "/" + screenshotName
        dropboxSSPath = dropboxPath + "/" + screenshotName
        #copy Screenshot to Dropbox
        copyfile(screenshotPath, dropboxSSPath)

         #******************LOOGGGING*************************
        #print("! (2/3) MARP Text Screenshot Copied to Dropbox")
        tempLog += "! (2/3) MARP Text Screenshot Copied to Dropbox" + "\n"

        #delete Screenshot
        os.remove(screenshotPath)

        #******************LOOGGGING*************************
        #print("! (3/3) MARP Text Screenshot Deleted")
        tempLog += "! (3/3) MARP Text Screenshot Deleted" + "\n"

    except:
        #print("* Error Copying MARP Files to Dropbox")
        tempLog += "* Error Copying MARP Files to Dropbox" + "\n"
    
    return tempLog

#...Screenshot Methods
def get_latest_image(dirpath, valid_extensions=('jpeg')):
    """
    Get the latest image file in the given directory
    """

    # get filepaths of all files and dirs in the given dir
    valid_files = [os.path.join(dirpath, filename) for filename in os.listdir(dirpath)]

  

    # filter out directories, no-extension, and wrong extension files
    valid_files = [f for f in valid_files if '.' in f and \
        f.rsplit('.',1)[-1] in valid_extensions and os.path.isfile(f)]

    if not valid_files:
        raise ValueError("No valid images in %s" % dirpath)

    return max(valid_files, key=os.path.getmtime)
def getScreenshotName():

    #get last name of member
    memberLastName = str(argPerson).split()[1]
    #edit date for filename
    editedDate = argDateScreen.replace("/", "-")
    #edit time for filename
    editedTime = str(time_converter(argResponseTime)).replace(":", "_")
    editedTime = editedTime.replace(".", "")
    editedTime = editedTime.replace(" ", "")
    #create new filename
    screenshotFilename = memberLastName + "_" + editedDate + "(" + editedTime + ").jpeg"

    return screenshotFilename
def renameScreenshot():


    #create a tempLog to return
    tempLog = ""  

    try:

        #Get most recent image screenshot
        strLatestImage = get_latest_image(argScreenshotDIR)

             

        #******************LOOGGGING*************************
        #print("! (1/2) MARP Text Screenshot Loaded")
        tempLog += "! (1/2) MARP Text Screenshot Loaded" + "\n"

        fileName = getScreenshotName()
        screenshotPath = argScreenshotDIR + "/" + fileName

        #rename file
        os.rename(strLatestImage, screenshotPath)

        #******************LOOGGGING*************************
        #print("! (2/2) MARP Text Screenshot Renamed")
        tempLog += "! (2/2) MARP Text Screenshot Renamed" + "\n"

    except:
        #print("* Error Renaming MARP Text Screenshot")
        tempLog += "* Error Renaming MARP Text Screenshot" + "\n"
    
    return tempLog

#PYTHONKIT Implementation
def runToolbox(_globWordTemplate, _argExcel_Filename, _argWord_FileDIR, _argNumPerMonth, _argDateScreen,
    _argInitialContactTime, _argResponseTime, _argPerson, _argTaskType, _argMonthYear, _argScreenshotDIR):

    #Assign Global Variables Through Swift Code

    #Connect with Global Variables
    global globWordTemplate
    global argExcel_Filename
    global argWord_FileDIR
    global argNumPerMonth
    global argDateScreen
    global argInitialContactTime
    global argResponseTime
    global argPerson
    global argTaskType
    global argMonthYear
    global argScreenshotDIR

    globWordTemplate = _globWordTemplate
    argExcel_Filename = _argExcel_Filename
    argWord_FileDIR = _argWord_FileDIR
    argNumPerMonth = _argNumPerMonth
    argDateScreen = _argDateScreen
    argInitialContactTime = _argInitialContactTime
    argResponseTime = _argResponseTime
    argPerson = _argPerson
    argTaskType = _argTaskType
    argMonthYear = _argMonthYear
    argScreenshotDIR = _argScreenshotDIR


    #FOR DEBUG:
    #argumentString = globWordTemplate + " " + argExcel_Filename + " " + argWord_FileDIR + " " + argNumPerMonth + " " + argDateScreen + " " + argInitialContactTime + " " + argResponseTime + " " + argPerson + " " +  argTaskType + " " + argMonthYear + " " + argScreenshotDIR
    #print(argumentString)
    #runToolbox("/Users/davidaustin/Google Drive/Programming/Workspace/Python/MARP Toolbox TESTING/Data/", "/Users/davidaustin/Google Drive/Programming/Workspace/Python/MARP Toolbox TESTING/Data/MARP Contact Template.docx", "/Users/davidaustin/Dropbox/MARP/DRUG SCREENS/Logs/Monthly Reports/May 2021/DrugScreen-CallList-(May-2021).xlsx", "/Users/davidaustin/Documents/MARP_Toolbox/Data/", "1", "10/9/1990", "10:05 AM", "10:30", "Millie Dyson", "SCREENSHOT", "May 2021", "/Users/davidaustin/Downloads/")
    #************

    #Create a run log
    processLog = ""

    if str(argTaskType).upper() == "EXCEL":

        result = LogContactDate()
        processLog += result

    elif str(argTaskType).upper() == "WORD":

        result = EditWordDocument()
        processLog += result

    elif str(argTaskType).upper() == "SCREENSHOT":

        result = renameScreenshot()
        processLog += result

    elif str(argTaskType).upper() == "COPY":

        result = CopyDocs()
        processLog += result

    return processLog





